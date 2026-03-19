"""解析路由器 - 根据文档类型和特征智能分发到不同解析器"""
import os
from typing import Dict, Any, Optional, Tuple
from parsers.utils.unified_loader import UnifiedLoader
from parsers.unstructured.unstructured_parser import UnstructuredParser
from parsers.parser_factory import ParserFactory
from utils.logger import logger


class ParsingRouter:
    """解析路由器 - 智能分发文档到合适的解析器"""
    
    # 解析器类型枚举
    PARSER_TYPE_UNSTRUCTURED = "unstructured"  # Unstructured 解析器（复杂格式、布局分析）
    PARSER_TYPE_LEGACY = "legacy"  # 原有解析器（标准格式）
    
    def __init__(self):
        """初始化路由器"""
        self.loader = UnifiedLoader()
        self.unstructured_parser = None
    
    def _get_unstructured_parser(self) -> Optional[UnstructuredParser]:
        """获取 Unstructured 解析器（延迟初始化）"""
        if self.unstructured_parser is None:
            try:
                self.unstructured_parser = UnstructuredParser()
            except Exception as e:
                logger.warning(f"Unstructured 解析器初始化失败: {e}")
                return None
        return self.unstructured_parser

    def _unstructured_pdf_available(self) -> bool:
        """
        检测 Unstructured 是否具备 PDF 分区能力。
        仅安装 `unstructured` 本体时，`partition_pdf()` 可能不可用，需要额外依赖（unstructured[pdf]）。
        """
        try:
            # 只要能导入 partition_pdf，通常就具备必要依赖
            from unstructured.partition.pdf import partition_pdf  # noqa: F401
            return True
        except Exception:
            return False
    
    def _detect_scanned_pdf(self, file_path: str) -> bool:
        """
        检测是否为扫描版PDF（需要OCR）
        
        Args:
            file_path: PDF文件路径
        
        Returns:
            是否为扫描版PDF
        """
        try:
            try:
                import PyPDF2
            except ImportError:
                logger.warning("PyPDF2未安装，无法检测扫描版PDF")
                return False
            
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                total_pages = len(pdf_reader.pages)
                
                if total_pages == 0:
                    return False
                
                # 检查前3页（或全部页面，如果少于3页）
                pages_to_check = min(3, total_pages)
                text_length = 0
                
                for i in range(pages_to_check):
                    try:
                        page_text = pdf_reader.pages[i].extract_text()
                        if page_text:
                            text_length += len(page_text.strip())
                    except Exception:
                        pass
                
                # 如果前几页几乎没有文本，可能是扫描版
                avg_text_per_page = text_length / pages_to_check
                if avg_text_per_page < 50:  # 每页少于50个字符
                    logger.info(f"检测到扫描版PDF: {file_path} (平均每页文本: {avg_text_per_page:.1f} 字符)")
                    return True
        except Exception as e:
            logger.warning(f"检测扫描版PDF时出错: {e}")
        
        return False
    
    def _detect_complex_format(self, file_path: str) -> bool:
        """
        检测是否为复杂格式文档（需要布局分析）
        
        检测标准：
        1. 扫描版PDF（需要OCR和布局分析）
        2. 包含大量图片的文档
        3. 包含复杂表格的文档
        4. 包含公式的文档
        5. 大文件（可能包含复杂布局）
        
        Args:
            file_path: 文件路径
        
        Returns:
            是否为复杂格式文档
        """
        file_ext = os.path.splitext(file_path)[1].lower()
        
        # 1. 检测扫描版PDF
        if file_ext == '.pdf':
            if self._detect_scanned_pdf(file_path):
                return True
        
        # 2. 检查文件大小（大文件更可能是复杂格式）
        try:
            file_size = os.path.getsize(file_path)
            # 如果文件大于2MB，优先使用Unstructured
            if file_size > 2 * 1024 * 1024:
                logger.info(f"检测到大型文档: {file_path} (大小: {file_size / 1024 / 1024:.2f} MB)")
                return True
        except Exception as e:
            logger.warning(f"获取文件大小时出错: {e}")
        
        # 3. 对于Word文档，检查是否包含复杂元素
        if file_ext == '.docx':
            try:
                try:
                    from docx import Document
                except ImportError:
                    logger.warning("python-docx未安装，无法检查Word文档特征")
                    return False
                
                doc = Document(file_path)
                
                # 检查表格数量
                table_count = len(doc.tables)
                if table_count > 3:  # 超过3个表格
                    logger.info(f"检测到复杂格式 Word 文档（包含 {table_count} 个表格）: {file_path}")
                    return True
                
                # 检查图片数量（通过关系）
                image_count = 0
                try:
                    for rel in doc.part.rels.values():
                        if "image" in rel.target_ref.lower():
                            image_count += 1
                    if image_count > 5:  # 超过5张图片
                        logger.info(f"检测到复杂格式 Word 文档（包含 {image_count} 张图片）: {file_path}")
                        return True
                except Exception:
                    pass
                
                # 检查段落数量（大量段落可能表示复杂结构）
                if len(doc.paragraphs) > 100:
                    logger.info(f"检测到复杂格式 Word 文档（包含 {len(doc.paragraphs)} 个段落）: {file_path}")
                    return True
            except Exception as e:
                logger.warning(f"检查Word文档特征时出错: {e}")
        
        # 4. 对于PDF，尝试快速检测是否包含图片
        if file_ext == '.pdf':
            try:
                try:
                    import PyPDF2
                except ImportError:
                    logger.warning("PyPDF2未安装，无法检查PDF特征")
                    return False
                
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    # 检查前几页是否有文本
                    pages_to_check = min(3, len(pdf_reader.pages))
                    has_text = False
                    for i in range(pages_to_check):
                        try:
                            page_text = pdf_reader.pages[i].extract_text()
                            if page_text and len(page_text.strip()) > 100:
                                has_text = True
                                break
                        except Exception:
                            pass
                    
                    # 如果前几页没有文本，可能是扫描版或图片PDF
                    if not has_text and pages_to_check > 0:
                        logger.info(f"检测到可能的图片PDF或扫描版PDF: {file_path}")
                        return True
            except Exception as e:
                logger.warning(f"检查PDF特征时出错: {e}")
        
        return False
    
    def _should_use_legacy_parser(self, file_path: str) -> bool:
        """
        判断是否应该使用原有解析器
        
        原有解析器适合：
        1. 简单的文本版PDF
        2. 简单的Word文档（无复杂表格、图片）
        3. Markdown、TXT等简单格式
        
        Args:
            file_path: 文件路径
        
        Returns:
            是否使用原有解析器
        """
        file_ext = os.path.splitext(file_path)[1].lower()
        
        # 简单格式文件直接使用原有解析器
        if file_ext in ['.txt', '.md', '.markdown']:
            return True
        
        # 对于PDF和Word，如果检测到复杂格式，不使用原有解析器
        if file_ext in ['.pdf', '.docx', '.doc']:
            if self._detect_complex_format(file_path):
                return False
        
        # 默认使用原有解析器
        return True
    
    def route(self, file_path: str) -> Tuple[str, Any]:
        """
        路由文档到合适的解析器
        
        路由策略：
        1. 复杂格式文档（扫描版PDF、包含大量图片/表格的文档）-> Unstructured解析器
        2. 简单格式文档（文本版PDF、简单Word、Markdown、TXT）-> 原有解析器
        
        Args:
            file_path: 文件路径
        
        Returns:
            (解析器类型, 解析器实例) 元组
        """
        # 验证文件
        if not self.loader.validate_file(file_path):
            raise ValueError(f"无效的文件: {file_path}")
        
        file_ext = os.path.splitext(file_path)[1].lower()
        
        # 0. 显式路由：仅 Unstructured 支持的格式 -> 直接走 Unstructured
        UNSTRUCTURED_ONLY_EXTENSIONS = ('.pptx', '.xlsx', '.xls', '.html', '.htm')
        if file_ext in UNSTRUCTURED_ONLY_EXTENSIONS:
            parser = self._get_unstructured_parser()
            if parser:
                logger.info(f"✓ 显式路由到 Unstructured 解析器: {file_path} (扩展名: {file_ext})")
                return self.PARSER_TYPE_UNSTRUCTURED, parser
            logger.warning(f"Unstructured 解析器不可用，将尝试 ParserFactory: {file_path}")
        
        # 1. 检测复杂格式文档 -> Unstructured解析器
        if self._detect_complex_format(file_path):
            # 对于 PDF：如果 unstructured[pdf] 未安装，则不要路由到 Unstructured，避免先报错再回退造成噪音日志
            if file_ext == ".pdf" and not self._unstructured_pdf_available():
                logger.warning(
                    f"检测到复杂 PDF 但 Unstructured 缺少 PDF 依赖（建议安装 unstructured[pdf]），将直接使用原有解析器: {file_path}"
                )
            else:
                parser = self._get_unstructured_parser()
                if parser:
                    logger.info(f"✓ 路由到 Unstructured 解析器: {file_path} (复杂格式)")
                    return self.PARSER_TYPE_UNSTRUCTURED, parser
                else:
                    logger.warning(f"Unstructured解析器不可用，回退到原有解析器: {file_path}")
        
        # 2. 使用原有解析器（适合简单格式）
        parser = ParserFactory.get_parser(file_path)
        if parser:
            logger.info(f"✓ 路由到原有解析器: {file_path} (标准格式)")
            return self.PARSER_TYPE_LEGACY, parser
        
        raise ValueError(f"无法找到合适的解析器: {file_path}")

